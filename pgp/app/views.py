from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from django.utils import timezone
from django.http import HttpResponse
from datetime import timedelta
from datetime import datetime
from .models import Round, Song, Player, Vote, PlayerStats
from .forms import RoundForm, SongSubmissionForm, DynamicVoteForm, UserProfileForm
from .spotify_views import get_spotify_client
from .spotify_utils import resolve_spotify_url, validate_spotify_url
from urllib.parse import urlparse
from docx.shared import Pt
from django.db.models import Sum, Value
from django.db.models.functions import Coalesce
import re
from django.contrib.auth import logout
from docx import Document
import logging
from django.contrib.auth import login
from django.contrib.auth.forms import UserCreationForm
from .utils import get_combined_song_data, get_round_winners, LoggedInPlayerStats, get_user_chart_data
from django.core.paginator import Paginator, EmptyPage
from django.http import JsonResponse
import json
from django.core.cache import cache

import logging
logger = logging.getLogger(__name__)


def user_logout(request):
    logout(request)
    messages.success(request, "Du er no logga ut. Snakkast!")
    return redirect('home')

@login_required
def home(request):
    import time
    start_total = time.time()
    
    start = time.time()
    players = Player.objects.all()
    logger.info(f"[TIMING] Players query: {(time.time()-start)*1000:.1f}ms")
    
    start = time.time()
    active_rounds = Round.objects.filter(end_date__gte=timezone.now())
    logger.info(f"[TIMING] Active rounds: {(time.time()-start)*1000:.1f}ms")
    
    start = time.time()
    stats = PlayerStats.objects.all().order_by('-total_points')
    logger.info(f"[TIMING] Stats query: {(time.time()-start)*1000:.1f}ms")
    
    logged_in_player_stats = None
    chart_data = None
    
    if request.user.is_authenticated:
        try:
            start = time.time()
            player_stats = LoggedInPlayerStats(request.user)
            player = request.user.player
            logger.info(f"[TIMING] LoggedInPlayerStats init: {(time.time()-start)*1000:.1f}ms")

            start = time.time()
            logged_in_player_stats = {
                'previous_songs': player_stats.previous_songs().order_by('-id'),
                'top_voters': player_stats.top_voters(),
                'top_given_votes': player_stats.top_given_votes(),
                'top_score_12_songs': player_stats.top_score_12_songs(),
            }
            logger.info(f"[TIMING] All player stats: {(time.time()-start)*1000:.1f}ms")
            
            start = time.time()
            chart_data = get_user_chart_data(player)
            logger.info(f"[TIMING] Chart data: {(time.time()-start)*1000:.1f}ms")
            
        except Player.DoesNotExist:
            logged_in_player_stats = None
            chart_data = None

    start = time.time()
    result = render(request, 'app/home.html', {
        'players': players,
        'active_rounds': active_rounds,
        'stats': stats,
        'logged_in_player_stats': logged_in_player_stats,
        'chart_data': json.dumps(chart_data) if chart_data else None,
    })
    logger.info(f"[TIMING] Template render: {(time.time()-start)*1000:.1f}ms")
    logger.info(f"[TIMING] TOTAL VIEW TIME: {(time.time()-start_total)*1000:.1f}ms")
    
    return result




@login_required
def user_profile(request):
    """Display user profile with edit functionality"""
    try:
        player = request.user.player
    except:
        player = None
    
    if request.method == 'POST':
        form = UserProfileForm(request.POST, instance=request.user)
        
        if form.is_valid():
            form.save()
            messages.success(request, '🎉 Profilen din er oppdatert!')
            return redirect('user_profile')
        else:
            messages.error(request, '❌ Det oppstod ein feil. Sjekk informasjonen og prøv igjen.')
    else:
        form = UserProfileForm(instance=request.user)

    # Get organizer rounds
    organizer_rounds = Round.objects.filter(organizer=request.user).order_by('-start_date')

    context = {
        'form': form,
        'player': player,
        'organizer_rounds': organizer_rounds,
    }
    return render(request, 'registration/profile.html', context)



@login_required
def round_list(request):
    rounds = Round.objects.all().order_by('-start_date')
    return render(request, 'app/round_list.html', {'rounds': rounds, 'now': timezone.now()})

@login_required
def create_round(request):
    if request.method == 'POST':
        form = RoundForm(request.POST)
        if form.is_valid():
            start_date_str = request.POST.get('start_date')
            end_date_str = request.POST.get('end_date')
            start_time_str = request.POST.get('start_time', '00:00')
            end_time_str = request.POST.get('end_time', '00:00')

            try:
                start_datetime = datetime.combine(
                    datetime.strptime(start_date_str, '%Y-%m-%d').date(),
                    datetime.strptime(start_time_str, '%H:%M').time()
                )
                end_datetime = datetime.combine(
                    datetime.strptime(end_date_str, '%Y-%m-%d').date(),
                    datetime.strptime(end_time_str, '%H:%M').time()
                )
            except ValueError:
                messages.error(request, 'Invalid date or time format.')
                return render(request, 'app/create_round.html', {'form': form})

            round = form.save(commit=False)
            round.organizer = request.user
            round.start_date = start_datetime
            round.end_date = end_datetime
            round.save()

            return redirect('round_detail', pk=round.pk)
    else:
        form = RoundForm()
    return render(request, 'app/create_round.html', {'form': form})

def get_track_id(player_song):
    """
    Extracts track ID from a player's song.
    Handles both old and new Spotify URL formats.
    """
    if player_song and player_song.spotify_url:
        # Try the new URL resolution method
        from .spotify_utils import extract_track_id
        track_id = extract_track_id(player_song.spotify_url)
        if track_id:
            return track_id
        
        # Fallback to the old method for backward compatibility
        parsed_url = urlparse(player_song.spotify_url)
        path_parts = parsed_url.path.split('/')
        if len(path_parts) > 2 and path_parts[1] == 'track':
            return path_parts[2]
    
    return None

def calculate_countdown(now, release_time):
    time_since_release = now - release_time
    if time_since_release < timedelta(hours=2):
        countdown = timedelta(hours=2) - time_since_release
        total_seconds = int(countdown.total_seconds())
        return f"{total_seconds // 3600}:{(total_seconds // 60) % 60:02}:{total_seconds % 60:02}"
    return None


@login_required
def round_detail(request, pk):
    round_instance = get_object_or_404(Round, pk=pk)
    player = Player.objects.get(user=request.user)
    player_song = Song.objects.filter(round=round_instance, player=player).first()
    has_voted = Vote.objects.filter(player=player, song__round=round_instance).exists()
    now = timezone.now()
    form = SongSubmissionForm(instance=player_song)
    track_id = get_track_id(player_song)
    is_organizer = round_instance.organizer == request.user
    release_time = round_instance.release_time if round_instance.playlist_released else None
    countdown = calculate_countdown(now, release_time) if release_time else None

    # Calculate total scores for songs in this round
    all_submitted_songs = Song.objects.filter(round=round_instance).annotate(
        aggregated_score=Coalesce(Sum('vote__score'), Value(0))  # Default to 0 if no votes
    ).order_by('-aggregated_score')

    submitted_player_ids = all_submitted_songs.values_list('player_id', flat=True).distinct()

    player_votes = Vote.objects.filter(player=player, song__round=round_instance) \
        .select_related('song') \
        .order_by('-score')  # Sort votes by score
    
     # Fetch all players in the database
    all_players = Player.objects.all()

    # Get the list of players who have voted
    players_with_votes_ids = set(Vote.objects.filter(song__round=round_instance).values_list('player_id', flat=True).distinct())

    # Get players who have submitted a song (these are already filtered players for this round)
    players_with_song = Player.objects.filter(id__in=submitted_player_ids)

    # Players who have not voted
    players_without_votes = players_with_song.exclude(id__in=players_with_votes_ids)

    # Fetch all players who have not submitted a song
    players_without_song = all_players.exclude(id__in=submitted_player_ids)

    has_submitted_song = player_song is not None

    if is_organizer and request.method == "POST" and 'finish_round' in request.POST:
        round_instance.round_finished = True
        round_instance.save()
        return redirect('round_detail', pk=round_instance.pk)    
    
    # Update the context to include all necessary data for both organizers and players
    context = {
        'round': round_instance,
        'round_finished': round_instance.round_finished, 
        'form': form,
        'player_song': player_song,
        'track_id': track_id,
        'is_organizer': is_organizer,
        'now': now,
        'has_voted': has_voted,
        'release_time': release_time,
        'countdown': countdown,
        'all_submitted_songs': all_submitted_songs,
        'num_submitted_songs': all_submitted_songs.count(),
        'players': Player.objects.filter(id__in=submitted_player_ids), 
        'players_with_votes': Player.objects.filter(id__in=players_with_votes_ids),  # Players who have voted
        'players_without_votes': players_without_votes,  
        'players_without_song': players_without_song,
        'players_with_song': players_with_song,
        'all_players': all_players, 
        'votes_dict': {
            song.id: {
                vote.player.id: vote.score for vote in song.vote_set.all()
            }
            for song in all_submitted_songs
        },
        'player_votes': player_votes, 
        'has_submitted_song': has_submitted_song, 
    }

    return render(request, 'app/round_detail.html', context)



@login_required
def update_playlist_url(request, pk):
    round = get_object_or_404(Round, pk=pk)
    if request.method == 'POST':
        playlist_url = request.POST.get('playlist_url')

        # Validate that the URL is a valid Spotify playlist URL
        if playlist_url and not re.match(r'https?://(open\.spotify\.com|spotify\.com)/playlist/.+', playlist_url):
            messages.error(request, 'Forbanna tosk! Dette er ikkje ei Spotify-liste!')
            return redirect('round_detail', pk=pk)

        # If valid, update the playlist URL
        round.playlist_url = playlist_url
        if playlist_url:
            if not round.playlist_released:  # Set release_time only when first released
                round.release_time = timezone.now()  # Set the release time to current time
            round.playlist_released = True
            messages.success(request, 'pGP-liste publisert!')
        else:
            round.playlist_released = False
            round.release_time = None  # Reset release time if the playlist is removed
            messages.success(request, 'Spotify-liste fjernet!')

        round.save()
        return redirect('round_detail', pk=pk)
    
def remove_playlist(request, pk):
    round = get_object_or_404(Round, pk=pk)
    if request.user == round.organizer:
        round.playlist_url = None
        round.playlist_released = False
        round.save()
    return redirect('round_detail', pk=pk)

@login_required
def submit_song(request, pk):
    round = get_object_or_404(Round, pk=pk)
    player = Player.objects.get(user=request.user)
    player_song = Song.objects.filter(round=round, player=player).first()

    # Check if the current date and time is before the round's start_date
    if timezone.now() > round.start_date:
        messages.error(request, 'Ikke registrert - Fristen har passert!')
        return redirect('round_detail', pk=pk)

    if request.method == 'POST':
        form = SongSubmissionForm(request.POST, instance=player_song)
        
        # Get the raw Spotify URL from the form
        spotify_url_input = request.POST.get('spotify_url', '').strip()
        
        # Basic validation
        if not spotify_url_input:
            messages.error(request, 'Du må oppgi ei Spotify-lenke!')
            return render(request, 'app/round_detail.html', {
                'round': round,
                'form': form,
                'player_song': player_song,
            })
        
        # Validate it looks like a Spotify URL
        if not validate_spotify_url(spotify_url_input):
            messages.error(request, 'Dette ser ikkje ut som ei Spotify-lenke. Prøv igjen!')
            return render(request, 'app/round_detail.html', {
                'round': round,
                'form': form,
                'player_song': player_song,
            })
        
        # Resolve the URL (handles both spotify.link and open.spotify.com formats)
        standardized_url, track_id = resolve_spotify_url(spotify_url_input)
        
        if not standardized_url or not track_id:
            messages.error(request, 'Kunne ikkje tolke Spotify-lenka di. Prøv igjen!')
            return render(request, 'app/round_detail.html', {
                'round': round,
                'form': form,
                'player_song': player_song,
            })
        
        # Get Spotify client and fetch track data
        sp = get_spotify_client()
        
        if sp:
            try:
                track_data = sp.track(track_id)
                artist_name = track_data['artists'][0]['name']
                song_title = track_data['name']

                # Check if the song already exists (same artist and title)
                existing_song = Song.objects.filter(
                    round=round,
                    title=song_title,
                    artist=artist_name
                ).exclude(player=player).exists()

                if existing_song:
                    messages.error(request, '🖐 Stopp! Nokon har allereie levert denne! Prøy igjen med anna låt...')
                    return redirect('round_detail', pk=pk)
                
                # If the checks pass, update or create the song with standardized URL
                Song.objects.update_or_create(
                    round=round,
                    player=player,
                    defaults={
                        'spotify_url': standardized_url,
                        'title': song_title,
                        'artist': artist_name
                    }
                )

                messages.success(request, 'Bidrag registrert!')
                return redirect('round_detail', pk=pk)

            except Exception as e:
                logger.error(f"Spotify API error: {str(e)}")
                messages.error(request, 'Hmm? Vis dette til Jostein: {}'.format(str(e)))
        else:
            messages.error(request, 'Får ikkje kontakt med Spotify.')
    else:
        form = SongSubmissionForm(instance=player_song)

    return render(request, 'app/round_detail.html', {
        'round': round,
        'form': form,
        'player_song': player_song,
    })


# View to delete the player's song
def delete_song(request, pk):
    round_obj = get_object_or_404(Round, pk=pk)
    player_song = Song.objects.filter(round=round_obj, player=request.user.player).first()

    if player_song:
        player_song.delete()
        messages.success(request, 'Bidraget ditt ble slettet.')
    else:
        messages.error(request, 'Du har ikke noe bidrag å slette.')

    return redirect('round_detail', pk=pk)


logger = logging.getLogger(__name__)

@login_required
def vote_view(request, pk):
    round_instance = get_object_or_404(Round, pk=pk)
    player = Player.objects.get(user=request.user)
    player_song = Song.objects.filter(round=round_instance, player=player).first()
    songs = Song.objects.filter(round=round_instance).exclude(player=player).order_by('artist', 'title')

    if request.method == 'POST':
        form = DynamicVoteForm(request.POST, songs=songs)
        logger.debug("Form data: %s", request.POST)
        
        if form.is_valid():
            cleaned_data = form.cleaned_data
            logger.debug("Cleaned data: %s", cleaned_data)
            
            for field_name, score in cleaned_data.items():
                if score:
                    song_id = field_name.split('-')[1]
                    song = get_object_or_404(Song, pk=song_id)
                    
                    # Check if the player has already voted for this song
                    if not Vote.objects.filter(player=player, song=song).exists():
                        Vote.objects.create(player=player, song=song, score=int(score))
                        logger.debug(f"Vote saved: Player {player.nickname}, Song {song.title}, Score {score}")
                    else:
                        logger.warning(f"Duplicate vote attempt: Player {player.nickname}, Song {song.title}, Score {score}")
            
            messages.success(request, '🤩 Poenga dine er registrert!!')
            return redirect('round_detail', pk=pk)
        else:
            logger.error("Form is invalid")
            logger.debug("Form errors: %s", form.errors)
    else:
        form = DynamicVoteForm(songs=songs)

    return render(request, 'app/vote_form.html', {
        'round': round_instance,
        'form': form,
        'songs': songs,
        'player_song': player_song,
        'score_range': form.score_range,
    })

def export_to_word(request, round_id):
    # Get the relevant round
    round_instance = Round.objects.get(id=round_id)
    
    # Get all submitted songs for this round
    all_submitted_songs = Song.objects.filter(round=round_instance)

    # Get all players who submitted songs in the current round
    submitted_players = all_submitted_songs.values_list('player', flat=True).distinct()  # Get unique players who submitted songs
    
    # Get players who voted
    voted_players = Vote.objects.filter(song__round=round_instance).values_list('player', flat=True).distinct()
    voted_players_set = set(voted_players)  # Unique players who have voted

    # Identify disqualified players (those who haven't voted)
    disqualified_players = [player for player in submitted_players if player not in voted_players_set]

    # Create a new Document
    doc = Document()

    doc.add_heading(f"Oppsummering av runden: {round_instance.name}", level=1)

    # Add Disqualified players section
    if disqualified_players:
        doc.add_heading("🤢 Diska:", level=2)
        for player_id in disqualified_players:
            player_nickname = Player.objects.get(id=player_id).nickname  # Get the player's nickname
            
            # Get the song details for the player (assuming one song per player for simplicity)
            song = all_submitted_songs.filter(player_id=player_id).first()  # Get the first song submitted by the player
            if song:
                total_score = song.total_score
                artist = song.artist  # Adjust if you need to access the artist differently
                # Add the player nickname with song details in parentheses
                doc.add_paragraph(f"{player_nickname} ({artist} - {song.title}, {total_score} poeng)")

    # Add a blank line before the summary
    doc.add_paragraph("") 

    # Get only the songs that have been submitted by players who voted
    qualified_songs = all_submitted_songs.filter(player__in=voted_players_set)

    # Sort the songs by total_score in ascending order for ranking
    sorted_songs = qualified_songs.order_by('total_score')  # Change here to ascending order

    # Initialize rank variables
    total_songs = sorted_songs.count()
    previous_score = None
    current_rank = total_songs  # Start ranking from the highest score

    for index, song in enumerate(sorted_songs):
        total_score = song.total_score
        player_nickname = song.player.nickname  # The player who submitted the song
        
        # Check if the score has changed to adjust the rank
        if total_score != previous_score:
            current_rank = total_songs - index  # Adjust rank based on total songs minus index
        
        # Write song details to the document with rank
        header_paragraph = doc.add_paragraph()
        run = header_paragraph.add_run(f"{current_rank}. {song.artist} - {song.title} ({total_score} poeng)")
        run.bold = True
        run.font.size = Pt(14)  # Set font size for title
        run.font.name = 'Arial'  # Change font family
        
        # Create a new paragraph for "Levert av"
        delivered_paragraph = doc.add_paragraph()
        delivered_run = delivered_paragraph.add_run(f"Levert av: {player_nickname}")
        delivered_run.bold = True
        delivered_run.font.size = Pt(12)  # Set font size for "Levert av"
        delivered_run.font.name = 'Arial'  # Change font family for consistency
        
        doc.add_paragraph("Poeng:")
        
        # Get the votes for the song and group them by score
        votes = Vote.objects.filter(song=song)  # Get all votes for the current song
        
        # Create a dictionary to group player nicknames by score
        score_dict = {}
        for vote in votes:
            score = vote.score
            player_nickname = vote.player.nickname
            if score not in score_dict:
                score_dict[score] = []
            score_dict[score].append(player_nickname)

        # Sort the scores in descending order and write votes to the document
        for score in sorted(score_dict.keys(), reverse=True):  # Sort scores in descending order
            players = score_dict[score]
            players_str = ', '.join(players)  # Join player nicknames with commas
            doc.add_paragraph(f"{score} poeng fra {players_str}")

        # Add a blank line between songs
        doc.add_paragraph("")  # This adds a line break

    # Prepare the response
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="{round_instance.name}_export.docx"'
    
    doc.save(response)  # Save the document to the response
    return response




def register(request):
    if request.method == 'POST':
        form = UserCreationForm(request.POST)
        password1 = request.POST.get('password1')
        password2 = request.POST.get('password2')
        
        if password1 != password2:
            messages.error(request, "Passorda er ikkje like!")
            return render(request, 'registration/register.html', {'form': form})

        if form.is_valid():
            user = form.save()
            login(request, user)  # Log the user in immediately after registration
            messages.success(request, "Velkommen ombord - du er no registrert!")
            return redirect('home')  # Redirect to a home or success page

    else:
        form = UserCreationForm()
    
    return render(request, 'registration/register.html', {'form': form})

@login_required
def my_rounds(request):
    # Fetch rounds that the user is the organizer of
    rounds = Round.objects.filter(organizer=request.user)

    return render(request, 'app/my_rounds.html', {
        'rounds': rounds,
    })

@login_required
def edit_round(request, round_id):
    round_instance = get_object_or_404(Round, id=round_id, organizer=request.user)

    if request.method == 'POST':
        # Handle deletion
        if 'delete_round' in request.POST:
            round_instance.delete()
            return redirect('round_list')  # Redirect to a view that lists rounds after deletion

        # Handle form submission for editing the round
        form = RoundForm(request.POST, instance=round_instance)
        if form.is_valid():
            form.save()
            return redirect('round_detail', pk=round_instance.id)  # Ensure to use 'pk' here
    else:
        form = RoundForm(instance=round_instance)

    return render(request, 'app/edit_round.html', {'form': form, 'round': round_instance})

@login_required
def combined_song_data_view(request):
    page = int(request.GET.get('page', 1))
    page_size = int(request.GET.get('pageSize', 10))
    search_value = request.GET.get('search[value]', '').strip()
    order_column = request.GET.get('orderColumn', 'dato')
    order_dir = request.GET.get('orderDir', 'asc')

    # ✅ SOLUTION: Cache as dictionaries, not custom objects
    cache_key = 'combined_songs_data_dicts'
    combined_songs_dicts = cache.get(cache_key)
    
    if combined_songs_dicts is None:
        # Get the custom objects
        combined_songs_objects = get_combined_song_data()
        
        # Convert to dictionaries for caching (these CAN be pickled)
        combined_songs_dicts = [
            {
                'dato': song.dato,
                'artist': song.artist,
                'tittel': song.tittel,
                'levert_av': song.levert_av,
                'tema': song.tema,
                'spotify': song.spotify,
            }
            for song in combined_songs_objects
        ]
        
        cache.set(cache_key, combined_songs_dicts, 300)  # Cache for 5 minutes
        
    total_records = len(combined_songs_dicts)

    # Early exit for short search terms
    if search_value and len(search_value) < 3:
        search_value = ''

    # Filter dictionaries instead of objects
    if search_value:
        search_lower = search_value.lower()
        
        filtered_songs = []
        for song_dict in combined_songs_dicts:
            if (search_lower in song_dict['artist'].lower() or
                search_lower in song_dict['tittel'].lower() or
                search_lower in song_dict['levert_av'].lower() or
                search_lower in song_dict['tema'].lower()):
                filtered_songs.append(song_dict)
        
        combined_songs_dicts = filtered_songs
    
    total_filtered = len(combined_songs_dicts)

    # Sort dictionaries
    if order_column in ['dato', 'artist', 'tittel', 'levert_av', 'tema']:
        reverse_sort = (order_dir == 'desc')
        combined_songs_dicts.sort(key=lambda x: x[order_column], reverse=reverse_sort)

    # Slice for pagination
    start_idx = (page - 1) * page_size
    end_idx = start_idx + page_size
    songs_slice = combined_songs_dicts[start_idx:end_idx]

    # songs_slice is already in the right format!
    return JsonResponse({
        'recordsTotal': total_records,
        'recordsFiltered': total_filtered,
        'draw': int(request.GET.get('draw', 1)),
        'data': songs_slice  # Already dictionaries, ready to go
    })

@login_required
def combined_song_view(request):
    return render(request, 'app/combined_songs.html')

def get_winners_data():
    # Get the round winners data (this returns a list of namedtuples)
    winners = get_round_winners()

    # Prepare the winners data for the response
    return [
        {
            'round_name': winner.round_name,
            'organizer': winner.organizer,
            'date_added': winner.date_added,
            'artist': winner.artist,
            'song_title': winner.song_title,
            'spotify_url': winner.spotify_url,
            'winning_player_nickname': winner.winning_player_nickname,
        }
        for winner in winners
    ]

@login_required
def round_winners_data_view(request):
    # Call the helper function to get winners data
    winners_data = get_winners_data()

    # Return the data as JSON
    return JsonResponse({
        'data': winners_data
    })

@login_required
def round_winners_view(request):
    # Call the helper function to get winners data
    winners_data = get_winners_data()

    # Pass the data to the template context
    return render(request, 'app/round_winners.html', {'winners': winners_data})